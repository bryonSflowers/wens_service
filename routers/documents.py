import json
import io
from datetime import datetime

from fastapi import APIRouter, Depends, Form, HTTPException, UploadFile, File
from typing import Optional

import db as db_service
from middleware import get_current_user

router = APIRouter(prefix="/documents", tags=["Documents"])

ALLOWED = {
    ".txt": "text",
    ".csv": "csv",
    ".xlsx": "excel",
    ".xls": "excel",
    ".docx": "word",
    ".doc": "word",
}


def _parse_text(raw: bytes) -> str:
    return raw.decode("utf-8", errors="replace")


def _parse_csv(raw: bytes) -> tuple[str, list[list[str]]]:
    text = raw.decode("utf-8", errors="replace")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    table = [line.split(",") for line in lines]
    return text, table


def _parse_excel(raw: bytes) -> tuple[str, list[list[str]]]:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    parts: list[str] = []
    tables: list[list[list[str]]] = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"--- Sheet: {sheet} ---")
        sheet_rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            cleaned = [str(c) if c is not None else "" for c in row]
            sheet_rows.append(cleaned)
            parts.append("\t".join(cleaned))
        tables.append(sheet_rows)
    content = "\n".join(parts)
    return content, tables


def _parse_word(raw: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(raw))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
):
    ext = ("." + file.filename.rsplit(".", 1)[-1]).lower() if "." in (file.filename or "") else ""
    if ext not in ALLOWED:
        raise HTTPException(400, f"Unsupported file type '{ext}'. Allowed: {', '.join(ALLOWED)}")
    file_type = ALLOWED[ext]
    raw = await file.read()
    if len(raw) > 20 * 1024 * 1024:
        raise HTTPException(400, "File too large (max 20MB)")
    try:
        if file_type == "excel":
            content, tables = _parse_excel(raw)
            raw_tables = json.dumps(tables)
        elif file_type == "word":
            content = _parse_word(raw)
            raw_tables = None
        else:
            content = _parse_text(raw)
            raw_tables = None
    except Exception as e:
        raise HTTPException(400, f"Failed to parse file: {str(e)}")
    word_count = len(content.split())
    pool = await db_service.get_pool()
    row = await pool.fetchrow(
        "INSERT INTO uploaded_docs (user_id, filename, file_type, content, raw_tables, word_count, metadata) "
        "VALUES ($1, $2, $3, $4, $5::jsonb, $6, $7::jsonb) RETURNING *",
        current_user["id"], file.filename, file_type, content, raw_tables, word_count,
        json.dumps({"uploaded_at": datetime.utcnow().isoformat()}),
    )
    await db_service.log_audit(pool, current_user["id"], "document.upload", "uploaded_docs", str(row["id"]))
    return db_service._serialize_row(row)


@router.get("")
async def list_documents(
    page: int = 1,
    page_size: int = 20,
    current_user: dict = Depends(get_current_user),
):
    pool = await db_service.get_pool()
    return await db_service.get_paginated(
        pool, "uploaded_docs",
        where="user_id = $1",
        params=[current_user["id"]],
        order_by="created_at DESC",
        page=page, page_size=page_size,
    )


@router.get("/{doc_id}")
async def get_document(doc_id: int, current_user: dict = Depends(get_current_user)):
    pool = await db_service.get_pool()
    row = await pool.fetchrow(
        "SELECT * FROM uploaded_docs WHERE id = $1 AND user_id = $2", doc_id, current_user["id"]
    )
    if not row:
        raise HTTPException(404, "Document not found")
    return db_service._serialize_row(row)


@router.delete("/{doc_id}", status_code=204)
async def delete_document(doc_id: int, current_user: dict = Depends(get_current_user)):
    pool = await db_service.get_pool()
    r = await pool.execute("DELETE FROM uploaded_docs WHERE id = $1 AND user_id = $2", doc_id, current_user["id"])
    if r == "DELETE 0":
        raise HTTPException(404, "Document not found")
